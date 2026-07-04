#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
医工学実験レポート docx ビルダー。
下地慶英の過去レポートの構成・書式を再現する。

使い方:
    python build_report.py <spec.json> <out.docx> [template.docx]
template を省略すると同梱の assets/template.docx を使う。

必要: python-docx (pip install python-docx)

spec(JSON) 形式:
{
  "cover": {
     "series":   "医工学実験Ⅱ　レポート",     # 24pt 中央 ゴシック太字
     "subtitle": "実験題目　　...",           # 16pt 左  ゴシック 下線
     "date":     "実験日　　2026年 ... 日",    # 16pt 中央 下線
     "submit":   "提出日　　2026年 ... 日",    # 16pt 中央 下線
     "affil":    "医工学コース　学生証番号　　　24TB4039",  # 16pt 左 下線
     "name":     "氏名　　　下地　慶英",         # 16pt 左 下線
     "collaborators": ["共同実験者　　...", ...]           # 16pt 左 下線
  },
  "body": [ <block>, ... ]
}

block の type:
  h1 / h2 / h3 : 見出し   {text}          -> Heading 1/2/3 (16/14/12pt 太字)
  p            : 本文段落 {text}          -> 段落頭を全角字下げ
  b            : 太字の小見出し行 {text}   -> 字下げなし太字(原理章の「① AND（論理積）」等)
  p_red        : 赤文字段落 {text}        -> ユーザーが実験後に埋める箇所
  ul           : 箇条書き {items:[...]}   -> List Paragraph
  ul_red       : 赤箇条書き {items:[...]}
  eq           : 数式行 {text, num}       -> 中央に数式、右端に式番号(例 num="式 2-1")
                 否定(バー)は各文字に結合文字U+0305を付けて近似する(例 "A̅・̅B̅")。
                 厳密な表記が必要ならWordの数式で置き換えるよう赤注記を添える。
  figure       : 図(未挿入) {ph, caption} -> 赤プレースホルダ(画像位置) + 下にキャプション(黒中央)
  fig_done     : 図(挿入済) {path, caption, width_in?} -> 画像 + 下にキャプション
  table        : 表(未挿入) {caption, ph} -> キャプション(黒中央,上) + 下に赤プレースホルダ
  table_data   : 表(実データ) {caption, header:[...], rows:[[...]], group_header?:[["入力",N],["出力",M]]}
                 -> キャプション(上) + 三線表(横罫線のみ・中央)。本人のテキスト表は三線表。
  code         : コード/実行結果 {text, lead?} -> lead("〜を以下に示す。")の後に等幅・罫線枠(1セル表)。
                 プログラミング系(C1〜C5)でソース・実行コマンド・出力を貼るときに使う。
  spacer       : 空行

キャプションの番号は「図 4-1　題名」「表 2-1　題名」のように
図/表と番号の間に半角スペース、番号と題名の間に全角スペースを入れる。
"""
import sys, json, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn

MINCHO = "ＭＳ 明朝"
GOTHIC = "ＭＳ ゴシック"
RED = RGBColor(0xFF, 0x00, 0x00)


def set_font(run, name=MINCHO, size=None, bold=None, color=None, underline=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = rpr.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rf)
    for k in ('w:ascii', 'w:hAnsi', 'w:eastAsia'):
        rf.set(qn(k), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if underline is not None:
        run.font.underline = underline


def clear_body(doc):
    body = doc.element.body
    for child in list(body):
        if child.tag == qn('w:sectPr'):
            continue
        body.remove(child)


def add_para(doc, text="", style=None, align=None, red=False, mincho=True,
             size=None, bold=None, underline=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_font(run, name=(MINCHO if mincho else GOTHIC), size=size,
                 bold=bold, color=(RED if red else None), underline=underline)
    return p


def build_cover(doc, c):
    """表紙。実物は項目間に空白が多く、値の行に下線が付く。"""
    def gap(n=1):
        for _ in range(n):
            add_para(doc)
    add_para(doc, c["series"], align=WD_ALIGN_PARAGRAPH.CENTER, mincho=False,
             size=24, bold=True)
    gap(2)
    add_para(doc, c.get("subtitle", ""), align=WD_ALIGN_PARAGRAPH.LEFT,
             mincho=False, size=16, underline=True)
    gap(2)
    add_para(doc, c.get("date", ""), align=WD_ALIGN_PARAGRAPH.CENTER,
             mincho=False, size=16, underline=True)
    gap(1)
    add_para(doc, c.get("submit", ""), align=WD_ALIGN_PARAGRAPH.CENTER,
             mincho=False, size=16, underline=True)
    gap(2)
    add_para(doc, c.get("affil", ""), align=WD_ALIGN_PARAGRAPH.CENTER,
             mincho=False, size=16, underline=True)
    add_para(doc, c.get("name", ""), align=WD_ALIGN_PARAGRAPH.CENTER,
             mincho=False, size=16, underline=True)
    gap(2)
    for co in c.get("collaborators", []):
        add_para(doc, co, align=WD_ALIGN_PARAGRAPH.CENTER, mincho=False,
                 size=16, underline=True)
    doc.add_page_break()


def _cell_border(cell, top=None, bottom=None):
    """セルの上下罫線のみ設定する(三線表用)。値は線の太さ(1/8pt)。"""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = tcPr.makeelement(qn('w:tcBorders'), {})
        tcPr.append(borders)
    for edge, sz in (("top", top), ("bottom", bottom)):
        if sz is None:
            continue
        el = borders.find(qn('w:' + edge))
        if el is None:
            el = borders.makeelement(qn('w:' + edge), {})
            borders.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:color'), '000000')


def _fill_cell(cell, text, bold=False, red=False):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    set_font(run, bold=bold, color=(RED if red else None))


def build_block(doc, b):
    t = b["type"]
    if t in ("h1", "h2", "h3"):
        add_para(doc, b["text"], style="Heading " + t[1])
    elif t == "p":
        add_para(doc, "　" + b["text"])
    elif t == "b":
        add_para(doc, b["text"], bold=True)
    elif t == "p_red":
        add_para(doc, "　" + b["text"], red=True)
    elif t == "ul":
        for it in b["items"]:
            add_para(doc, it, style="List Paragraph")
    elif t == "ul_red":
        for it in b["items"]:
            add_para(doc, it, style="List Paragraph", red=True)
    elif t == "eq":
        # 数式は中央、式番号(例「式 2-1」)は右端。タブストップで実現する。
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.tab_stops.add_tab_stop(Cm(7.5), WD_TAB_ALIGNMENT.CENTER)
        pf.tab_stops.add_tab_stop(Cm(15.0), WD_TAB_ALIGNMENT.RIGHT)
        num = b.get("num", "")
        run = p.add_run("\t" + b["text"] + ("\t" + num if num else ""))
        set_font(run)
    elif t == "figure":
        add_para(doc, b["ph"], align=WD_ALIGN_PARAGRAPH.CENTER, red=True)
        add_para(doc, b["caption"], style="Caption",
                 align=WD_ALIGN_PARAGRAPH.CENTER)
    elif t == "fig_done":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(b["path"], width=Inches(b.get("width_in", 5.0)))
        add_para(doc, b["caption"], style="Caption",
                 align=WD_ALIGN_PARAGRAPH.CENTER)
    elif t == "table":
        add_para(doc, b["caption"], style="Caption",
                 align=WD_ALIGN_PARAGRAPH.CENTER)
        add_para(doc, b["ph"], align=WD_ALIGN_PARAGRAPH.CENTER, red=True)
    elif t == "table_data":
        add_para(doc, b["caption"], style="Caption",
                 align=WD_ALIGN_PARAGRAPH.CENTER)
        header = b.get("header")
        rows = b["rows"]
        group = b.get("group_header")  # 例 [["入力", 3], ["出力", 1]]
        ncols = len(header) if header else len(rows[0])
        tbl = doc.add_table(rows=0, cols=ncols)
        tbl.alignment = 1  # 中央
        # グループヘッダ行(入力/出力)。セルを結合して中央に見出しを置く。
        if group:
            cells = tbl.add_row().cells
            idx = 0
            for label, span in group:
                first = cells[idx]
                last = cells[idx + span - 1]
                merged = first.merge(last) if span > 1 else first
                _fill_cell(merged, label)
                idx += span
            for c in tbl.rows[0].cells:
                _cell_border(c, top=12)
        if header:
            cells = tbl.add_row().cells
            for c, txt in zip(cells, header):
                _fill_cell(c, txt)
            for c in cells:
                if not group:
                    _cell_border(c, top=12, bottom=6)
                else:
                    _cell_border(c, bottom=6)
        for row in rows:
            cells = tbl.add_row().cells
            for c, txt in zip(cells, row):
                _fill_cell(c, txt)
        for c in tbl.rows[-1].cells:
            _cell_border(c, bottom=12)
    elif t == "code":
        # プログラミング系: 「〜を以下に示す。」の後に等幅・罫線枠(1セル表)でコード/出力を貼る
        lead = b.get("lead")
        if lead:
            add_para(doc, "　" + lead)
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = 1
        cell = tbl.rows[0].cells[0]
        _cell_border(cell, top=6, bottom=6)
        tcPr = cell._tc.get_or_add_tcPr()
        borders = tcPr.find(qn('w:tcBorders'))
        for edge in ("left", "right"):
            el = borders.makeelement(qn('w:' + edge), {})
            el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '6'); el.set(qn('w:color'), '000000')
            borders.append(el)
        p0 = cell.paragraphs[0]
        first = True
        for line in b["text"].split("\n"):
            p = p0 if first else cell.add_paragraph()
            first = False
            run = p.add_run(line if line else "")
            set_font(run, name="Menlo", size=9)
    elif t == "spacer":
        add_para(doc)
    else:
        raise ValueError("unknown block type: " + t)


def build(spec, out, template=None, overwrite=False):
    if template is None:
        template = os.path.join(os.path.dirname(__file__), "..", "assets",
                                "template.docx")
    # 安全ガード: 引数順ミスや既存レポートの破壊を防ぐ
    if os.path.abspath(out) == os.path.abspath(template):
        raise ValueError("出力先とテンプレートが同一。引数順を確認せよ: build(spec, out, template)")
    if os.path.exists(out) and not overwrite:
        raise FileExistsError(
            f"出力先が既に存在する: {out}\n"
            "既存ファイルを上書きする場合は overwrite=True (CLI: --overwrite) を指定する。"
            "ユーザーの完成レポートを上書きしないこと。")
    doc = Document(template)
    clear_body(doc)
    build_cover(doc, spec["cover"])
    for b in spec["body"]:
        build_block(doc, b)
    doc.save(out)
    return out


if __name__ == "__main__":
    args = [a for a in sys.argv[1:] if a != "--overwrite"]
    ow = "--overwrite" in sys.argv
    if len(args) < 2:
        print(__doc__)
        sys.exit(1)
    spec = json.load(open(args[0], encoding="utf-8"))
    tmpl = args[2] if len(args) > 2 else None
    print("wrote", build(spec, args[1], tmpl, overwrite=ow))
